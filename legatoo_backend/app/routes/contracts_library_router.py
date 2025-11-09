"""
Contracts Library API Router

RESTful endpoints for managing contracts, templates, revisions, and AI generation.
"""

from typing import Optional
from types import SimpleNamespace
from fastapi import APIRouter, Depends, HTTPException, status, Query
from sqlalchemy.ext.asyncio import AsyncSession

from ..db.database import get_db
from ..utils.auth import get_current_user
from ..schemas.profile_schemas import TokenData
from ..schemas.contracts_library import (
    ContractCreate, ContractUpdate, ContractResponse, ContractListResponse,
    TemplateCreate, TemplateUpdate, TemplateResponse,
    RevisionCreate, RevisionHistoryResponse,
    AIGenerateRequest, AIGenerateResponse,
    ContractFilters, TemplateFilters
)
from ..services.contracts.contracts_library_service import ContractsLibraryService
from ..utils.contracts_formatting import (
    normalize_contract_content,
    parse_contract_structure,
    extract_plain_text,
)

router = APIRouter(prefix="/contracts", tags=["contracts"])


# ============ Contract Endpoints ============

@router.post("", response_model=ContractResponse, status_code=status.HTTP_201_CREATED)
async def create_contract(
    contract_data: ContractCreate,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Create a new contract."""
    service = ContractsLibraryService(db)
    try:
        contract = await service.create_contract(contract_data, current_user.sub)
        return contract
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.get("", response_model=ContractListResponse)
async def list_contracts(
    category: Optional[str] = Query(None),
    jurisdiction: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    language: Optional[str] = Query(None),
    ai_generated: Optional[bool] = Query(None),
    search_query: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """List contracts with filtering and pagination."""
    filters = ContractFilters(
        category=category,
        jurisdiction=jurisdiction,
        status=status,
        language=language,
        ai_generated=ai_generated,
        search_query=search_query,
        page=page,
        page_size=page_size
    )
    service = ContractsLibraryService(db)
    try:
        result = await service.list_contracts(filters, current_user.sub)
        return result
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.get("/{contract_id}", response_model=ContractResponse)
async def get_contract(
    contract_id: str,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Get contract by ID."""
    service = ContractsLibraryService(db)
    try:
        # Ensure user_id is an int (handle both int and UUID types from TokenData)
        user_id = int(current_user.sub) if isinstance(current_user.sub, (int, str)) and str(current_user.sub).isdigit() else None
        
        contract = await service.get_contract(contract_id, user_id)
        if not contract:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND, 
                detail=f"Contract with ID '{contract_id}' not found"
            )
        return contract
    except ValueError as e:
        if "Access denied" in str(e):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(e))
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.put("/{contract_id}", response_model=ContractResponse)
async def update_contract(
    contract_id: str,
    update_data: ContractUpdate,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Update contract."""
    service = ContractsLibraryService(db)
    try:
        contract = await service.update_contract(contract_id, update_data, current_user.sub)
        if not contract:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Contract not found")
        return contract
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.delete("/{contract_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_contract(
    contract_id: str,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Soft delete (archive) contract."""
    service = ContractsLibraryService(db)
    try:
        success = await service.delete_contract(contract_id, current_user.sub)
        if not success:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Contract not found")
        return None
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


# ============ AI Generation Endpoints ============

@router.post("/generate", response_model=AIGenerateResponse, status_code=status.HTTP_201_CREATED)
async def generate_contract_with_ai(
    request: AIGenerateRequest,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Generate a contract using AI."""
    service = ContractsLibraryService(db)
    try:
        result = await service.generate_contract_with_ai(request, current_user.sub)
        return result
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.post("/generate/{request_id}/save", response_model=ContractResponse, status_code=status.HTTP_201_CREATED)
async def save_ai_generated_contract(
    request_id: str,
    contract_data: ContractCreate,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Save AI-generated content as a contract."""
    service = ContractsLibraryService(db)
    try:
        contract = await service.save_ai_generated_contract(request_id, contract_data, current_user.sub)
        return contract
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


# ============ Revision Endpoints ============

@router.get("/{contract_id}/history", response_model=RevisionHistoryResponse)
async def get_revision_history(
    contract_id: str,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Get revision history for a contract."""
    service = ContractsLibraryService(db)
    try:
        history = await service.get_revision_history(contract_id, current_user.sub)
        return history
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.post("/{contract_id}/revise", response_model=ContractResponse, status_code=status.HTTP_201_CREATED)
async def create_revision(
    contract_id: str,
    revision_data: RevisionCreate,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Create a new revision for a contract."""
    from ..schemas.contracts_library import RevisionResponse
    service = ContractsLibraryService(db)
    try:
        revision = await service.create_revision(contract_id, revision_data, current_user.sub)
        # Return updated contract
        contract = await service.get_contract(contract_id, current_user.sub)
        return contract
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


# ============ Template Endpoints ============

@router.post("/templates", response_model=TemplateResponse, status_code=status.HTTP_201_CREATED)
async def create_template(
    template_data: TemplateCreate,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Create a new contract template."""
    service = ContractsLibraryService(db)
    try:
        template = await service.create_template(template_data, current_user.sub)
        return template
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.get("/templates", response_model=dict)
async def list_templates(
    category: Optional[str] = Query(None),
    jurisdiction: Optional[str] = Query(None),
    language: Optional[str] = Query(None),
    is_public: Optional[bool] = Query(None),
    tags: Optional[str] = Query(None),  # Comma-separated tags
    search_query: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    current_user: Optional[TokenData] = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """List contract templates with filtering."""
    tag_list = tags.split(",") if tags else None
    filters = TemplateFilters(
        category=category,
        jurisdiction=jurisdiction,
        language=language,
        is_public=is_public,
        tags=tag_list,
        search_query=search_query,
        page=page,
        page_size=page_size
    )
    service = ContractsLibraryService(db)
    try:
        user_id = current_user.sub if current_user else None
        result = await service.list_templates(filters, user_id)
        return result
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.get("/templates/{template_id}", response_model=TemplateResponse)
async def get_template(
    template_id: str,
    current_user: Optional[TokenData] = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Get template by ID."""
    service = ContractsLibraryService(db)
    try:
        user_id = current_user.sub if current_user else None
        template = await service.get_template(template_id, user_id)
        if not template:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
        return template
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.put("/templates/{template_id}", response_model=TemplateResponse)
async def update_template(
    template_id: str,
    update_data: TemplateUpdate,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Update template."""
    service = ContractsLibraryService(db)
    try:
        template = await service.update_template(template_id, update_data, current_user.sub)
        if not template:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
        return template
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.delete("/templates/{template_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_template(
    template_id: str,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Delete template."""
    service = ContractsLibraryService(db)
    try:
        success = await service.delete_template(template_id, current_user.sub)
        if not success:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")
        return None
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.post("/templates/{template_id}/generate", response_model=ContractResponse, status_code=status.HTTP_201_CREATED)
async def generate_from_template(
    template_id: str,
    placeholder_data: dict,
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Generate a contract from a template by replacing placeholders."""
    service = ContractsLibraryService(db)
    try:
        contract = await service.generate_from_template(template_id, placeholder_data, current_user.sub)
        return contract
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


# ============ Export Endpoints ============

@router.get("/{contract_id}/export")
async def export_contract(
    contract_id: str,
    format: str = Query("pdf", regex="^(pdf|docx)$"),
    current_user: TokenData = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Export contract as PDF or Word document with Arabic support."""
    from fastapi.responses import StreamingResponse, Response
    from io import BytesIO
    from urllib.parse import quote
    import os
    from pathlib import Path
    
    service = ContractsLibraryService(db)
    
    try:
        user_id = int(current_user.sub) if isinstance(current_user.sub, (int, str)) and str(current_user.sub).isdigit() else None
        contract = await service.get_contract(contract_id, user_id)
        
        if not contract:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Contract not found"
            )
        
        if not contract.content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Contract has no content to export"
            )

        normalized_content = normalize_contract_content(contract.content)
        structured_elements = parse_contract_structure(normalized_content)
        plain_title = extract_plain_text(contract.title)
        plain_body = extract_plain_text(contract.content)

        def contains_rtl_characters(text: str) -> bool:
            if not text:
                return False
            for char in text:
                if "\u0590" <= char <= "\u08FF":
                    return True
            return False

        language_value = (contract.language or "").lower()
        rtl_prefixes = ("ar", "fa", "he", "ku", "ps", "ur", "sd", "dv")
        language_is_rtl = bool(
            language_value
            and any(language_value.startswith(prefix) for prefix in rtl_prefixes)
        )

        # Determine language (default to detecting RTL characters when language is missing or incorrect)
        is_arabic = language_is_rtl or contains_rtl_characters(plain_body)

        if format == "pdf":
            # Generate PDF
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import (
                    SimpleDocTemplate,
                    Paragraph,
                    Spacer,
                )
                from reportlab.lib.enums import TA_RIGHT, TA_LEFT
                from html import escape as html_escape

                buffer = BytesIO()

                # Create PDF with RTL-friendly margins
                doc = SimpleDocTemplate(
                    buffer,
                    pagesize=A4,
                    topMargin=0.75*inch,
                    bottomMargin=0.75*inch,
                    leftMargin=0.75*inch,
                    rightMargin=0.75*inch
                )
                styles = getSampleStyleSheet()

                # Register Arabic font
                arabic_font_name = None
                try:
                    from reportlab.pdfbase import pdfmetrics
                    from reportlab.pdfbase.ttfonts import TTFont

                    app_dir = Path(__file__).parent.parent
                    fonts_dir = app_dir / "fonts"

                    font_files = [
                        "Arial Unicode MS.ttf",
                        "Arial.ttf",
                        "NotoSansArabic-Regular.ttf",
                    ]

                    for font_file in font_files:
                        font_path = fonts_dir / font_file
                        if font_path.exists():
                            font_name = Path(font_file).stem.replace(" ", "")
                            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                            arabic_font_name = font_name
                            break
                    if not arabic_font_name:
                        system_font_candidates = [
                            Path("/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf"),
                            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
                            Path("/Library/Fonts/Arial Unicode.ttf"),
                            Path("/Library/Fonts/Arial Unicode MS.ttf"),
                            Path("C:/Windows/Fonts/arial.ttf"),
                            Path("C:/Windows/Fonts/arialuni.ttf"),
                            Path("C:/Windows/Fonts/tahoma.ttf"),
                        ]
                        for font_path in system_font_candidates:
                            if font_path.exists():
                                font_name = font_path.stem.replace(" ", "")
                                pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
                                arabic_font_name = font_name
                                break
                except Exception:
                    pass

                base_font = arabic_font_name if arabic_font_name else "Helvetica"
                alignment = TA_RIGHT if is_arabic else TA_LEFT
                word_wrap = "RTL" if is_arabic else None

                # Styles
                title_style = ParagraphStyle(
                    "ContractTitle",
                    parent=styles["Heading1"],
                    fontName=base_font,
                    fontSize=20,
                    leading=24,
                    textColor="#1f2937",
                    spaceBefore=6,
                    spaceAfter=16,
                    alignment=alignment,
                    wordWrap=word_wrap,
                )

                normal_style = ParagraphStyle(
                    "ContractNormal",
                    parent=styles["Normal"],
                    fontName=base_font,
                    fontSize=12,
                    leading=20,
                    textColor="#4b5563",
                    spaceAfter=12,
                    alignment=alignment,
                    wordWrap=word_wrap,
                )

                list_style = ParagraphStyle(
                    "ContractList",
                    parent=normal_style,
                    bulletFontName=base_font,
                    bulletFontSize=12,
                    bulletIndent=0,
                    leftIndent=0,
                    rightIndent=0,
                    spaceAfter=8,
                    alignment=alignment,
                    wordWrap=word_wrap,
                )

                # Process Arabic text
                try:
                    import arabic_reshaper
                    from bidi.algorithm import get_display
                    ARABIC_AVAILABLE = True
                except ImportError:
                    ARABIC_AVAILABLE = False

                def contains_arabic(text: str) -> bool:
                    return any("\u0600" <= char <= "\u06FF" for char in text)

                def rtl_wrap(text: str) -> str:
                    if not text:
                        return ""
                    if ARABIC_AVAILABLE and is_arabic and contains_arabic(text):
                        reshaped = arabic_reshaper.reshape(text)
                        bidi_text = get_display(reshaped)
                        return f"\u202B{bidi_text}\u202C"
                    return text

                def runs_to_html(runs):
                    output_parts = []
                    for run in runs or []:
                        raw_text = run.text or ""
                        raw_text = raw_text.replace("\r\n", "\n")
                        segments = raw_text.split("\n")
                        processed_segments = []
                        for segment in segments:
                            safe = rtl_wrap(segment)
                            safe = html_escape(safe)
                            processed_segments.append(safe)
                        joined = "<br/>".join(processed_segments)
                        if run.bold:
                            joined = f"<b>{joined}</b>"
                        if run.italic:
                            joined = f"<i>{joined}</i>"
                        if run.underline:
                            joined = f"<u>{joined}</u>"
                        output_parts.append(joined)
                    return "".join(output_parts) or "&nbsp;"

                # Build PDF content
                story = []
                story.append(Paragraph(runs_to_html([SimpleNamespace(text=plain_title, bold=True, italic=False, underline=False)]), title_style))
                story.append(Spacer(1, 0.15 * inch))

                metadata_lines = []
                if contract.category:
                    metadata_lines.append(
                        f"<b>{rtl_wrap('التصنيف' if is_arabic else 'Category')}:</b> "
                        f"{html_escape(rtl_wrap(contract.category))}"
                    )
                if contract.jurisdiction:
                    metadata_lines.append(
                        f"<b>{rtl_wrap('الاختصاص' if is_arabic else 'Jurisdiction')}:</b> "
                        f"{html_escape(rtl_wrap(contract.jurisdiction))}"
                    )
                if metadata_lines:
                    story.append(Paragraph("<br/>".join(metadata_lines), normal_style))
                    story.append(Spacer(1, 0.1 * inch))

                if not structured_elements:
                    story.append(Paragraph(runs_to_html([SimpleNamespace(text=extract_plain_text(contract.content), bold=False, italic=False, underline=False)]), normal_style))
                else:
                    for element in structured_elements:
                        if element.type == "heading":
                            heading_style = ParagraphStyle(
                                f"HeadingLevel{element.level}",
                                parent=title_style if element.level == 1 else normal_style,
                                fontName=base_font,
                                fontSize=20 - (element.level or 1) * 2,
                                leading=22,
                                alignment=alignment,
                                spaceBefore=12,
                                spaceAfter=10,
                                wordWrap=word_wrap,
                            )
                            story.append(Paragraph(runs_to_html(element.runs or []), heading_style))
                        elif element.type == "list" and element.items:
                            for idx, item_runs in enumerate(element.items, start=1):
                                bullet = f"{idx}." if element.ordered else "•"
                                bullet = rtl_wrap(bullet)
                                story.append(
                                    Paragraph(
                                        runs_to_html(item_runs),
                                        list_style,
                                        bulletText=bullet,
                                    )
                                )
                            story.append(Spacer(1, 0.08 * inch))
                        else:
                            story.append(Paragraph(runs_to_html(element.runs or []), normal_style))

                doc.build(story)
                buffer.seek(0)

                safe_title = "".join(c for c in contract.title if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_title = safe_title.encode('ascii', 'ignore').decode('ascii')
                filename = f"contract_{contract_id}_{safe_title}.pdf"
                encoded_filename = quote(filename, safe='')

                return StreamingResponse(
                    buffer,
                    media_type="application/pdf",
                    headers={
                        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
                    }
                )

            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="PDF generation requires reportlab library"
                )
        
        elif format == "docx":
            # Generate Word document
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn

                doc = Document()
                base_font_name = "Arial"

                # Set document direction to RTL
                if is_arabic:
                    for section in doc.sections:
                        section.page_height = Inches(11.69)  # A4
                        section.page_width = Inches(8.27)
                        section.left_margin = Inches(0.75)
                        section.right_margin = Inches(0.75)
                        section.top_margin = Inches(0.75)
                        section.bottom_margin = Inches(0.75)

                normal_style = doc.styles["Normal"]
                normal_style.font.name = base_font_name
                normal_style.font.size = Pt(12)
                normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), base_font_name)
                normal_style._element.rPr.rFonts.set(qn("w:cs"), base_font_name)

                def configure_paragraph(paragraph, align_right: bool = False, space_after: int = 12):
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT if (align_right or is_arabic) else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_after = Pt(space_after)
                    paragraph.paragraph_format.line_spacing = 1.8

                def add_runs(paragraph, runs, font_size: Optional[int] = None):
                    for run in runs or []:
                        text = (run.text or "").replace("\r\n", "\n")
                        segments = text.split("\n")
                        for idx, segment in enumerate(segments):
                            doc_run = paragraph.add_run(segment)
                            doc_run.font.name = base_font_name
                            doc_run.font.size = Pt(font_size if font_size else 12)
                            doc_run.font.bold = run.bold
                            doc_run.font.italic = run.italic
                            doc_run.font.underline = run.underline
                            rfonts = doc_run._element.rPr.rFonts
                            rfonts.set(qn("w:eastAsia"), base_font_name)
                            rfonts.set(qn("w:cs"), base_font_name)
                            if idx < len(segments) - 1:
                                doc_run.add_break()

                # Header with contract name
                section = doc.sections[0]
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = plain_title
                configure_paragraph(header_para, align_right=is_arabic, space_after=0)
                for run in header_para.runs:
                    run.font.name = base_font_name
                    run.font.size = Pt(11)
                    rfonts = run._element.rPr.rFonts
                    rfonts.set(qn("w:eastAsia"), base_font_name)
                    rfonts.set(qn("w:cs"), base_font_name)

                # Title
                title_paragraph = doc.add_paragraph()
                configure_paragraph(title_paragraph, align_right=is_arabic, space_after=18)
                add_runs(title_paragraph, [SimpleNamespace(text=plain_title, bold=True, italic=False, underline=False)], font_size=20)

                # Metadata
                if contract.category:
                    paragraph = doc.add_paragraph()
                    configure_paragraph(paragraph, align_right=is_arabic, space_after=8)
                    label = "التصنيف" if is_arabic else "Category"
                    add_runs(
                        paragraph,
                        [
                            SimpleNamespace(text=f"{label}: ", bold=True, italic=False, underline=False),
                            SimpleNamespace(text=contract.category, bold=False, italic=False, underline=False),
                        ],
                    )
                if contract.jurisdiction:
                    paragraph = doc.add_paragraph()
                    configure_paragraph(paragraph, align_right=is_arabic, space_after=12)
                    label = "الاختصاص" if is_arabic else "Jurisdiction"
                    add_runs(
                        paragraph,
                        [
                            SimpleNamespace(text=f"{label}: ", bold=True, italic=False, underline=False),
                            SimpleNamespace(text=contract.jurisdiction, bold=False, italic=False, underline=False),
                        ],
                    )

                # Content
                if not structured_elements:
                    paragraph = doc.add_paragraph()
                    configure_paragraph(paragraph, align_right=is_arabic, space_after=12)
                    add_runs(paragraph, [SimpleNamespace(text=extract_plain_text(contract.content), bold=False, italic=False, underline=False)])
                else:
                    size_map = {1: 18, 2: 16, 3: 15, 4: 14, 5: 13, 6: 12}
                    for element in structured_elements:
                        if element.type == "heading":
                            paragraph = doc.add_paragraph()
                            configure_paragraph(paragraph, align_right=is_arabic, space_after=12)
                            heading_size = size_map.get(element.level or 3, 14)
                            add_runs(paragraph, element.runs, font_size=heading_size)
                        elif element.type == "list" and element.items:
                            for idx, item_runs in enumerate(element.items, start=1):
                                paragraph = doc.add_paragraph()
                                configure_paragraph(paragraph, align_right=is_arabic, space_after=8)
                                bullet_prefix = rtl_wrap(f"{idx}. " if element.ordered else "• ")
                                prefix_run = SimpleNamespace(text=bullet_prefix, bold=True, italic=False, underline=False)
                                add_runs(paragraph, [prefix_run] + item_runs)
                        else:
                            paragraph = doc.add_paragraph()
                            configure_paragraph(paragraph, align_right=is_arabic, space_after=12)
                            add_runs(paragraph, element.runs)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                safe_title = "".join(c for c in contract.title if c.isalnum() or c in (' ', '-', '_')).strip()
                safe_title = safe_title.encode('ascii', 'ignore').decode('ascii')
                filename = f"contract_{contract_id}_{safe_title}.docx"
                encoded_filename = quote(filename, safe='')

                return StreamingResponse(
                    buffer,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
                    }
                )
                
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Word document generation requires python-docx library"
                )
        
    except HTTPException:
        raise
    except Exception as e:
        from ...config.enhanced_logging import get_logger
        logger = get_logger(__name__)
        logger.exception(f"Error exporting contract {contract_id}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to export contract: {str(e)}"
        )