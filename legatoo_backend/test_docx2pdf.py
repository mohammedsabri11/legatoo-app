#!/usr/bin/env python
"""Test script to diagnose docx2pdf conversion issues."""

import os
import sys
import tempfile
import traceback
from pathlib import Path

print("Testing docx2pdf conversion...")
print("=" * 60)

# Create a simple test DOCX file
try:
    from docx import Document
    
    print("Creating test DOCX file...")
    doc = Document()
    doc.add_paragraph("This is a test document for PDF conversion.")
    doc.add_paragraph("If you can read this, the conversion worked!")
    
    test_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(test_docx.name)
    test_docx.close()
    
    print(f"✅ Created test DOCX: {test_docx.name}")
    
    # Try conversion
    pdf_path = test_docx.name.replace(".docx", ".pdf")
    print(f"\nAttempting to convert to PDF: {pdf_path}")
    
    try:
        from docx2pdf import convert
        print("✅ docx2pdf imported successfully")
        
        print("\nCalling convert()...")
        convert(test_docx.name, pdf_path)
        
        if os.path.exists(pdf_path):
            file_size = os.path.getsize(pdf_path)
            print(f"✅ SUCCESS! PDF created at: {pdf_path}")
            print(f"   File size: {file_size} bytes")
        else:
            print(f"❌ FAILED: PDF file not found at: {pdf_path}")
            print("   Check if path is correct")
            
    except ImportError as e:
        print(f"❌ FAILED: Cannot import docx2pdf")
        print(f"   Error: {e}")
        print("   Install with: pip install docx2pdf")
        
    except Exception as e:
        print(f"❌ FAILED: Conversion error occurred")
        print(f"   Error Type: {type(e).__name__}")
        print(f"   Error Message: {str(e)}")
        print("\n   Full traceback:")
        traceback.print_exc()
        
        # Check for specific issues
        error_str = str(e).lower()
        if "libreoffice" in error_str or "soffice" in error_str:
            print("\n⚠️  Issue: LibreOffice not found")
            print("   Solution: Install LibreOffice from https://www.libreoffice.org/")
        elif "word" in error_str or "com" in error_str or "comtypes" in error_str:
            print("\n⚠️  Issue: Microsoft Word COM access failed")
            print("   Possible solutions:")
            print("   1. Ensure Microsoft Word is installed")
            print("   2. Close any open Word instances")
            print("   3. Restart Python/server")
            print("   4. Try installing LibreOffice instead (recommended)")
        elif "permission" in error_str or "access" in error_str:
            print("\n⚠️  Issue: Permission/Access denied")
            print("   Try running as administrator")
    
    finally:
        # Cleanup
        if os.path.exists(test_docx.name):
            try:
                os.unlink(test_docx.name)
                print(f"\n🧹 Cleaned up: {test_docx.name}")
            except:
                pass
        if os.path.exists(pdf_path):
            try:
                os.unlink(pdf_path)
                print(f"🧹 Cleaned up: {pdf_path}")
            except:
                pass
                
except ImportError:
    print("❌ FAILED: python-docx not installed")
    print("   Install with: pip install python-docx")
except Exception as e:
    print(f"❌ FAILED: {type(e).__name__}: {e}")
    traceback.print_exc()

print("\n" + "=" * 60)
print("Test complete!")
